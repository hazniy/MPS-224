from docx import Document # pip install python-docx
# You may need to modify this line to point to the directory where you have your Word files.
data_dir = '../data/office/'

def show_text(filename):
    """Print all the text in a Word file."""
    # Construct an instance of the class docx.Document, which represents a Word file.
    doc = Document(data_dir + filename)
    # The doc object has an attribute called paragraphs, as we could learn by reading 
    # the documentation or typing dir(doc).  The value of this attribute is a list of
    # instances of the class docx.text.paragraph.Paragraph.  Each of these instances
    # has an attribute called text, which is a string containing the text of the 
    # paragraph (without any of the information about fonts, colours and so on).
    text = ''
    for para in doc.paragraphs:
        text = text + para.text + '\n'
    print(text)
show_text('mouse.docx')

def find_paragraph(filename, s):
    """
    Looks through the specified file, and prints the first paragraph that 
    contains the string s.  If no such paragraph is found, prints 'Not found'.
    """
    doc = Document(data_dir + filename)
    for para in doc.paragraphs:
        if s.lower() in para.text.lower():
            return para.text
    return 'Not found'
print(find_paragraph('abc.docx','Zanzibar'))
